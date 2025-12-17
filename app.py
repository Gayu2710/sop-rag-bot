import streamlit as st
import chromadb
from groq import Groq
import time
import docx

st.set_page_config(page_title="SOP Assistant | AI", page_icon="🤖", layout="wide")

st.markdown("""<style>
.main-header{font-size:3rem;font-weight:bold;background:linear-gradient(90deg,#667eea 0%,#764ba2 100%);-webkit-background-clip:text;-webkit-text-color:transparent;}
.stat-box{background:#1e293b;padding:1.5rem;border-radius:10px;border-left:4px solid #667eea;margin:1rem 0}
</style>""", unsafe_allow_html=True)

@st.cache_resource
def init_chatbot():
    try:
        groq_client=Groq(api_key=st.secrets["GROQ_API_KEY"])
        client=chromadb.PersistentClient(path="./chromadb")
        collection=client.get_or_create_collection("sop_chunks")
        return groq_client,collection,client
    except Exception as e:
        st.error(f"❌ Initialization error: {e}")
        return None,None,None

groq_client,collection,client=init_chatbot()

def chunk_text(text:str,chunk_size:int=1000,chunk_overlap:int=200):
    chunks,start=[],0
    while start<len(text):
        chunks.append(text[start:start+chunk_size])
        start+=chunk_size-chunk_overlap
    return chunks

def process_docx(file):
    doc=docx.Document(file)
    full_text=[]
    for para in doc.paragraphs:
        if para.text.strip():
            full_text.append(para.text)
    for table in doc.tables:
        for row in table.rows:
            row_text=' | '.join([cell.text.strip() for cell in row.cells])
            if row_text.strip():
                full_text.append(row_text)
    return "\n".join(full_text)

def answer_sop(question):
    start=time.time()
    results=collection.query(query_texts=[question],n_results=5)
    ids=results["ids"][0]
    context="\n\n---\n\n".join(results["documents"][0])
    prompt=f"""You are an SOP assistant. Use ONLY the context below.

Context: {context}
Question: {question}
Answer:"""
    chat=groq_client.chat.completions.create(model="llama-3.3-70b-versatile",messages=[{"role":"user","content":prompt}],temperature=0.1)
    return chat.choices[0].message.content,ids[0],int((time.time()-start)*1000)

st.markdown('<h1 class="main-header">🤖 Incident Management SOP Assistant</h1>',unsafe_allow_html=True)
st.markdown('<p style="text-align:center;color:#94a3b8;font-size:1.1rem">AI-Powered Knowledge Base for DevOps & SRE Teams</p>',unsafe_allow_html=True)

if collection.count()==0:
    st.info("📋 **First-time setup:** Upload your SOP document")
    uploaded_file=st.file_uploader("📁 Upload SOP Document",type=['docx'])
    if uploaded_file:
        with st.spinner("⚡ Processing..."):
            try:
                raw_text=process_docx(uploaded_file)
                st.success(f"✅ Extracted **{len(raw_text):,}** characters")
                chunks=chunk_text(raw_text)
                st.info(f"📝 Created **{len(chunks)}** chunks")
                ids=[f"chunk-{i}" for i in range(len(chunks))]
                collection.upsert(ids=ids,documents=chunks,metadatas=[{"index":i} for i in range(len(chunks))])
                st.success(f"🎉 Indexed {len(chunks)} chunks! Refresh to start chatting.")
                st.balloons()
            except Exception as e:
                st.error(f"❌ Error: {e}")
else:
    col1,col2,col3=st.columns(3)
    with col1:
        st.markdown(f'<div class="stat-box"><h3>📊 Database</h3><p style="font-size:2rem;font-weight:bold;color:#10b981">{collection.count()}</p><p style="color:#94a3b8">Chunks Indexed</p></div>',unsafe_allow_html=True)
    with col2:
        st.markdown('<div class="stat-box"><h3>⚡ Speed</h3><p style="font-size:2rem;font-weight:bold;color:#3b82f6">~690ms</p><p style="color:#94a3b8">Avg Response</p></div>',unsafe_allow_html=True)
    with col3:
        st.markdown('<div class="stat-box"><h3>🤖 Model</h3><p style="font-size:1.3rem;font-weight:bold;color:#8b5cf6">Llama-3.3-70b</p><p style="color:#94a3b8">Groq API</p></div>',unsafe_allow_html=True)

    with st.sidebar:
        st.markdown("### 💡 Example Questions")
        for icon,q in [("🔴","What are the severity levels?"),("📅","Update cadence for Sev2?"),("⚠️","Handle NodeNotReady?"),("🔍","Tools for detection?"),("📝","What to document?")]:
            if st.button(f"{icon} {q}",key=q,use_container_width=True):
                st.session_state.clicked_question=q
        st.divider()
        st.markdown("### 🛠️ Tech Stack\n- **🗄️** ChromaDB\n- **🧠** Llama-3.3-70b\n- **🎨** Streamlit\n- **🔗** RAG Pipeline")
        st.divider()
        if st.button("🔄 Reset Database",type="secondary",use_container_width=True):
            client.delete_collection("sop_chunks")
            st.success("✅ Reset! Refresh page.")
                    st.cache_resource.clear()
            time.sleep(1)
            st.rerun()

    if "messages" not in st.session_state:
        st.session_state.messages=[]

    for msg in st.session_state.messages:
        with st.chat_message(msg["role"],avatar="🧑‍💻" if msg["role"]=="user" else "🤖"):
            st.markdown(msg["content"])
            if "metadata" in msg:
                st.caption(f"📎 {msg['metadata']['source']} | ⚡ {msg['metadata']['latency']}ms")

    prompt=st.session_state.pop("clicked_question",None) or st.chat_input("💬 Ask about incidents, procedures...")

    if prompt:
        st.session_state.messages.append({"role":"user","content":prompt})
        with st.chat_message("user",avatar="🧑‍💻"):
            st.markdown(prompt)
        with st.chat_message("assistant",avatar="🤖"):
            with st.spinner("🔍 Searching..."):
                response,source,latency=answer_sop(prompt)
                st.markdown(response)
                st.caption(f"📎 `{source}` | ⚡ {latency}ms")
        st.session_state.messages.append({"role":"assistant","content":response,"metadata":{"source":source,"latency":latency}})
